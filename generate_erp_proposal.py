#!/usr/bin/env python3
"""Generate ATA & LTA ERP Proposal in Speech-Enabled Microlearning Platform format."""

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn


def set_run_font(run, name='Calibri', size=11, bold=False, italic=False, color=None):
    font = run.font
    font.name = name
    font.size = Pt(size)
    font.bold = bold
    font.italic = italic
    if color:
        font.color.rgb = RGBColor(*color)
    run._element.rPr.rFonts.set(qn('w:eastAsia'), name)


def add_heading_paragraph(doc, text, bold=True, size=12, align=WD_ALIGN_PARAGRAPH.LEFT, space_after=6, color=None):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.space_before = Pt(0)
    run = p.add_run(text)
    set_run_font(run, size=size, bold=bold, color=color)
    return p


def add_body_paragraph(doc, text, indent=False, bold=False, italic=False, size=11, space_after=6, align=WD_ALIGN_PARAGRAPH.LEFT):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.space_before = Pt(0)
    if indent:
        p.paragraph_format.left_indent = Inches(0.25)
    run = p.add_run(text)
    set_run_font(run, size=size, bold=bold, italic=italic)
    return p


def add_bullet_paragraph(doc, text, indent_level=0, bold=False, size=11, space_after=4):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.left_indent = Inches(0.25 + indent_level * 0.25)
    run = p.add_run(text)
    set_run_font(run, size=size, bold=bold)
    return p


def add_plain_paragraph(doc, text, size=11, space_after=6, align=WD_ALIGN_PARAGRAPH.LEFT, bold=False):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.space_before = Pt(0)
    run = p.add_run(text)
    set_run_font(run, size=size, bold=bold)
    return p


def add_section_number_heading(doc, number, title, size=12):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.space_before = Pt(12)
    run = p.add_run(f"{number}) {title}")
    set_run_font(run, size=size, bold=True)
    return p


def main():
    doc = Document()

    # Set default font for the document
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)
    style._element.rPr.rFonts.set(qn('w:eastAsia'), 'Calibri')

    # ============================================================
    # HEADER (company info)
    # ============================================================
    add_plain_paragraph(doc, "MicroAxis", size=11, space_after=2)
    add_plain_paragraph(doc, "Contact Person: Mark Anthony Ureta", size=11, space_after=2)
    add_plain_paragraph(doc, "Email: mark@microaxis.ph", size=11, space_after=2)
    add_plain_paragraph(doc, "Phone: +63 XXX XXX XXXX", size=11, space_after=12)

    # ============================================================
    # TITLE
    # ============================================================
    add_heading_paragraph(doc, "FORMAL PROJECT PROPOSAL", size=14, bold=True, align=WD_ALIGN_PARAGRAPH.LEFT, space_after=12)

    add_heading_paragraph(doc, "Enterprise Resource Planning (ERP) System", size=13, bold=True, align=WD_ALIGN_PARAGRAPH.LEFT, space_after=4)
    add_plain_paragraph(doc, "A Web-Based System for Dual-Entity Accounting Firm Workflow, Billing, Disbursement, and Document Management", size=11, space_after=4)
    add_plain_paragraph(doc, "Prepared By: MicroAxis (\"Provider\")", size=11, space_after=2)
    add_plain_paragraph(doc, "Platform Type: Web-Based", size=11, space_after=2)
    add_plain_paragraph(doc, "Total Project Cost: PHP 37,000 (VAT exclusive)", size=11, space_after=18)

    # ============================================================
    # 1) EXECUTIVE SUMMARY
    # ============================================================
    add_section_number_heading(doc, "1", "Executive Summary")
    add_body_paragraph(doc,
        "MicroAxis proposes to design and develop a custom Enterprise Resource Planning (ERP) system for ATA & LTA Accounting Firm — a Philippine-based accounting practice operating under two distinct business entities: ATA (single proprietorship) and LTA (corporation). The proposed ERP will unify and streamline the firm’s core workflows from work request initiation through processing, billing, disbursement, and documentation management.",
        space_after=6)
    add_body_paragraph(doc,
        "Based on the provided workflow documentation and UI prototype, the system will include a unified login with role-based access, dual-entity client management, intelligent workflow and task tracking, billing with Philippine Peso (PHP) handling, disbursement approvals, and a digital document repository with original-copy handover tracking.",
        space_after=6)
    add_body_paragraph(doc,
        "The solution enables managerial staff to initiate and monitor work requests across six departments, empowers accounting and operations teams with structured task assignments and dependency tracking, and provides leadership with real-time visibility into revenue, outstanding balances, and firm performance.",
        space_after=12)

    # ============================================================
    # 2) PROJECT OBJECTIVES
    # ============================================================
    add_section_number_heading(doc, "2", "Project Objectives")
    add_bullet_paragraph(doc, "Deliver a role-based web platform for Admin, Manager, and Staff workflows across ATA and LTA entities.")
    add_bullet_paragraph(doc, "Implement dual-entity client management with proper data segregation, entity-specific dashboards, and client retainer tracking.")
    add_bullet_paragraph(doc, "Provide structured work request processing with task dependencies, assignment logic, and recurring retainer task automation.")
    add_bullet_paragraph(doc, "Enable billing for Professional Fees (PF) and government fees with automatic VAT calculations, PHP formatting, and aging reports.")
    add_bullet_paragraph(doc, "Build disbursement and expense tracking with configurable approval workflows (1-tier or 2-tier) and self-approval prevention.")
    add_bullet_paragraph(doc, "Deploy a Document Management System (DMS) with version control, categorization, and physical original copy handover logging.")
    add_bullet_paragraph(doc, "Generate managerial reports and analytics covering work request volume, task completion rates, billing summaries, and entity P&L snapshots.")
    add_bullet_paragraph(doc, "Support approximately 10 employees with audit trails and access controls aligned to Philippine data privacy principles.", space_after=12)

    # ============================================================
    # 3) SCOPE OF WORK
    # ============================================================
    add_section_number_heading(doc, "3", "Scope of Work")

    add_body_paragraph(doc, "A) Authentication & Role Access (Login)", bold=True, space_after=4)
    add_bullet_paragraph(doc, "Secure login (email/password) with session persistence", indent_level=1)
    add_bullet_paragraph(doc, "Role-based access control: Admin, Manager, and Staff", indent_level=1)
    add_bullet_paragraph(doc, "Entity-level restrictions: ATA-only, LTA-only, or cross-entity access", indent_level=1)
    add_bullet_paragraph(doc, "Entity switcher for multi-entity users with visual badge indicators", indent_level=1)

    add_body_paragraph(doc, "B) Multi-Entity Client Management Module", bold=True, space_after=4)
    add_bullet_paragraph(doc, "Client master records with mandatory entity assignment (ATA or LTA)", indent_level=1)
    add_bullet_paragraph(doc, "Retainer status tracking and recurring task template linkage", indent_level=1)
    add_bullet_paragraph(doc, "Search and filter by entity, service type, engagement status, and assigned staff", indent_level=1)
    add_bullet_paragraph(doc, "Client contact management and service history", indent_level=1)

    add_body_paragraph(doc, "C) Workflow & Task Management Module", bold=True, space_after=4)
    add_bullet_paragraph(doc, "Work request creation with title, description, client link, priority, and due date", indent_level=1)
    add_bullet_paragraph(doc, "Six-stage lifecycle: Pre-processing → Processing → Billing → Disbursement → Documentation → Completed", indent_level=1)
    add_bullet_paragraph(doc, "Task assignment to one or multiple employees with clear ownership", indent_level=1)
    add_bullet_paragraph(doc, "Task dependency mapping (DAG) — downstream tasks auto-start when upstream tasks are completed", indent_level=1)
    add_bullet_paragraph(doc, "Recurring retainer templates (e.g., monthly bookkeeping, quarterly tax filing)", indent_level=1)
    add_bullet_paragraph(doc, "Task status workflow: Draft → Assigned → In Progress → For Review → Completed → Cancelled", indent_level=1)
    add_bullet_paragraph(doc, "Dashboard views: My Tasks, Team Tasks, Overdue Tasks, Completed Today", indent_level=1)

    add_body_paragraph(doc, "D) Billing Module", bold=True, space_after=4)
    add_bullet_paragraph(doc, "Invoice generation per work request with line-item support for Professional Fees (PF) and government fees", indent_level=1)
    add_bullet_paragraph(doc, "Entity-specific invoice numbering and branding (ATA vs LTA)", indent_level=1)
    add_bullet_paragraph(doc, "PHP currency formatting with Philippine accounting standards (comma thousand separators, two decimal places)", indent_level=1)
    add_bullet_paragraph(doc, "Automatic VAT calculations with VATable and VAT-Exempt treatment", indent_level=1)
    add_bullet_paragraph(doc, "Invoice status tracking: Draft → Sent → Partially Paid → Paid → Overdue → Cancelled", indent_level=1)
    add_bullet_paragraph(doc, "Payment recording (method, date, reference number, amount) and aging reports (30/60/90+ days)", indent_level=1)
    add_bullet_paragraph(doc, "PDF export for client delivery", indent_level=1)

    add_body_paragraph(doc, "E) Disbursement & Expense Module", bold=True, space_after=4)
    add_bullet_paragraph(doc, "Employee expense filing with categories: Transportation, Notary, Meals, and Other", indent_level=1)
    add_bullet_paragraph(doc, "Approval workflow: Employee submits → Supervisor/Manager reviews → Accounting approves → Marked Released", indent_level=1)
    add_bullet_paragraph(doc, "Self-approval block: users cannot approve their own expense submissions", indent_level=1)
    add_bullet_paragraph(doc, "Client pass-through tracking: record government fees collected from client and paid to the agency", indent_level=1)
    add_bullet_paragraph(doc, "Fund source distinction: Firm Fund vs Client Fund", indent_level=1)
    add_bullet_paragraph(doc, "Receipt/attachment upload for every expense line item", indent_level=1)
    add_bullet_paragraph(doc, "Reimbursement summary by employee, department, and period", indent_level=1)

    add_body_paragraph(doc, "F) Document Management System (DMS)", bold=True, space_after=4)
    add_bullet_paragraph(doc, "File upload per work request with categorization (Requirement Docs, Processed Forms, Government Receipts, Final Deliverables, etc.)", indent_level=1)
    add_bullet_paragraph(doc, "Support for PDF, DOCX, XLSX, image formats, and scanned documents", indent_level=1)
    add_bullet_paragraph(doc, "Original copy tracking flag — indicates whether the physical original was handed to the client", indent_level=1)
    add_bullet_paragraph(doc, "Handover log with recipient name, date, and method (In-Person, Courier, Pickup)", indent_level=1)
    add_bullet_paragraph(doc, "Document version control with upload history and uploader attribution", indent_level=1)
    add_bullet_paragraph(doc, "Access restricted by user role and entity", indent_level=1)

    add_body_paragraph(doc, "G) Reporting & Analytics (MVP-lite)", bold=True, space_after=4)
    add_bullet_paragraph(doc, "Work Request Volume Report — count of requests by entity, department, and status", indent_level=1)
    add_bullet_paragraph(doc, "Task Completion Rate — average time from assignment to completion", indent_level=1)
    add_bullet_paragraph(doc, "Billing Summary — total PF and government fees billed, collected, and outstanding", indent_level=1)
    add_bullet_paragraph(doc, "Disbursement Report — total expenses by category and employee", indent_level=1)
    add_bullet_paragraph(doc, "Overdue Tasks Alert — automatic flagging of tasks past due date", indent_level=1)
    add_bullet_paragraph(doc, "Entity P&L Snapshot — revenue (billed) vs disbursements per entity (ATA vs LTA)", indent_level=1, space_after=12)

    # ============================================================
    # 4) DELIVERABLES
    # ============================================================
    add_section_number_heading(doc, "4", "Deliverables")
    add_bullet_paragraph(doc, "UI/UX implementation aligned to approved workflow and screens")
    add_bullet_paragraph(doc, "Complete web system: Dashboard, Clients, Workflow, Billing, Disbursement, Documents, Reports, and Admin portals")
    add_bullet_paragraph(doc, "Multi-entity client management with retainer tracking")
    add_bullet_paragraph(doc, "Workflow engine with task dependency mapping and retainer templates")
    add_bullet_paragraph(doc, "Billing module with PF/government fee line items, VAT, and aging reports")
    add_bullet_paragraph(doc, "Disbursement module with approval workflows and receipt upload")
    add_bullet_paragraph(doc, "Document management with version control and original-copy handover logging")
    add_bullet_paragraph(doc, "Reporting dashboards and analytics (MVP-lite)")
    add_bullet_paragraph(doc, "QA testing, UAT support, deployment, and basic documentation/orientation", space_after=12)

    # ============================================================
    # 5) ASSUMPTIONS / EXCLUSIONS
    # ============================================================
    add_section_number_heading(doc, "5", "Assumptions / Exclusions")
    add_bullet_paragraph(doc, "Third-party API integrations (BIR eFPS, eBIRForms, SEC eFAST, bank APIs, GCash, PayMongo) are excluded from MVP scope and may be addressed in Phase 2.")
    add_bullet_paragraph(doc, "Mobile native application (iOS/Android) is excluded; the MVP will be web-responsive only.")
    add_bullet_paragraph(doc, "Advanced OCR for scanned document full-text search is excluded from MVP.")
    add_bullet_paragraph(doc, "Multi-currency support beyond PHP and payroll module integration are excluded.")
    add_bullet_paragraph(doc, "External API integrations with third-party accounting software (e.g., QuickBooks, Xero) are excluded.")
    add_bullet_paragraph(doc, "Any new major modules beyond the listed navigation will be handled via change request.", space_after=12)

    # ============================================================
    # 6) PROJECT TIMELINE (10 to 12 Weeks)
    # ============================================================
    add_section_number_heading(doc, "6", "Project Timeline (10 to 12 Weeks)")
    add_body_paragraph(doc,
        "The project will be delivered within an estimated 10 to 12 weeks (approximately 2 to 3 months), depending on approval turnaround, content readiness, and UAT participation.",
        space_after=6)

    add_bullet_paragraph(doc, "Month 1 (Weeks 1–4): Milestone 1 — Discovery + Core Engine", bold=True)
    add_bullet_paragraph(doc, "Discovery & planning, UI/UX wireframes and prototype, authentication, roles, multi-entity client management, workflow engine baseline", indent_level=1)

    add_bullet_paragraph(doc, "Month 2 (Weeks 5–8): Milestone 2 — Business Modules", bold=True)
    add_bullet_paragraph(doc, "Billing module, disbursement module, document management system, reporting dashboards, module integration", indent_level=1)

    add_bullet_paragraph(doc, "Month 3 (Weeks 9–12): Milestone 3 — Testing, UAT, Deployment", bold=True)
    add_bullet_paragraph(doc, "Integration testing, QA, client UAT, bug fixes, production deployment, training sessions, and 14-day hypercare", indent_level=1, space_after=6)

    add_body_paragraph(doc,
        "Note: With fast approvals and complete inputs, delivery may be closer to 10 weeks; with more iterations, up to 12 weeks.",
        italic=True, space_after=12)

    # ============================================================
    # 7) PROJECT INVESTMENT
    # ============================================================
    add_section_number_heading(doc, "7", "Project Investment")
    add_plain_paragraph(doc, "Total Project Cost: PHP 37,000 (VAT exclusive)", size=11, space_after=12)

    # ============================================================
    # 8) PAYMENT TERMS
    # ============================================================
    add_section_number_heading(doc, "8", "Payment Terms")

    add_body_paragraph(doc, "Standard Milestone-Based Schedule:", bold=True, space_after=4)
    add_bullet_paragraph(doc, "40% upon project kickoff and contract signing: PHP 14,800")
    add_bullet_paragraph(doc, "30% upon completion of Sprint 1 (Alpha release, Week 5): PHP 11,100")
    add_bullet_paragraph(doc, "30% upon UAT sign-off and production deployment (Week 10–12): PHP 11,100")
    add_plain_paragraph(doc, "Total: 40% + 30% + 30% = 100% (PHP 37,000)", size=11, space_after=12)

    # ============================================================
    # 9) ACCEPTANCE, CHANGE REQUESTS, AND CONFIDENTIALITY
    # ============================================================
    add_section_number_heading(doc, "9", "Acceptance, Change Requests, and Confidentiality")
    add_bullet_paragraph(doc, "Each milestone is subject to Client review and written acceptance based on agreed deliverables.")
    add_bullet_paragraph(doc, "Scope additions or major revisions will follow a documented change request (time/cost impact provided prior to implementation).")
    add_bullet_paragraph(doc, "All client data, financial records, and proprietary business information will be treated as confidential.", space_after=12)

    # ============================================================
    # 10) CONCLUSION
    # ============================================================
    add_section_number_heading(doc, "10", "Conclusion")
    add_body_paragraph(doc,
        "This Enterprise Resource Planning (ERP) system will provide ATA & LTA Accounting Firm with a unified, role-based platform that streamlines work request workflows, automates billing and disbursement tracking, and maintains a secure digital document repository with physical handover logging. The proposed phased development ensures a practical build sequence and clear delivery checkpoints, with structured payment terms aligned to milestone completion.",
        space_after=18)

    # ============================================================
    # SIGN-OFF
    # ============================================================
    add_plain_paragraph(doc, "MicroAxis", size=11, bold=True, space_after=2)
    add_plain_paragraph(doc, "Business Development / Dept.", size=11, space_after=2)

    # Save
    output_path = "/home/javvii/FreelanceProject/Project4/ATA_LTA_ERP_Revised_Proposal.docx"
    doc.save(output_path)
    print(f"Saved: {output_path}")


if __name__ == "__main__":
    main()
